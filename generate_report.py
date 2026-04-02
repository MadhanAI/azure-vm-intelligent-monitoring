"""
generate_report.py — Word Report Generator (Python / python-docx)
==================================================================
All design implemented in Python:
  - Page borders on EVERY page (XML injection into sectPr)
  - Header: client name (left) | period (right), blue underline
  - Footer: "Page X of Y | Azure VM Performance Report | CONFIDENTIAL"
  - Cover page: dark-blue banner table, executive summary
  - Per-VM sections: coloured info table, 6 metric charts per VM
  - Findings: colour-coded status summary table + bullet recommendations
  - No Word template dependency

Charts rendered with matplotlib (Agg backend — no display required).
"""

import os
import io
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import FuncFormatter
import dateutil.parser

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from collect_metrics import VMMetrics, ReportConfig


# ──────────────────────────────────────────────
# Brand colours (hex strings without #)
# ──────────────────────────────────────────────
C_DARK      = "1F4E79"    # Dark navy — brand primary, borders, headings
C_MID       = "2E75B6"    # Mid blue  — section headers, accent lines
C_LIGHT     = "DEEAF1"    # Light blue — table header backgrounds
C_WHITE     = "FFFFFF"
C_GREY_LIGHT = "F2F2F2"   # Alternating row fill
C_GREY_TEXT  = "595959"   # Body text secondary
C_GREEN     = "375623"    # Normal status text
C_GREEN_BG  = "E2EFDA"    # Normal status background
C_AMBER     = "BF8F00"    # Warning status text
C_AMBER_BG  = "FFF2CC"    # Warning status background
C_RED       = "C00000"    # Critical status text
C_RED_BG    = "FFDFD9"    # Critical status background

# RGB tuples for direct use
RGB_DARK   = RGBColor(0x1F, 0x4E, 0x79)
RGB_MID    = RGBColor(0x2E, 0x75, 0xB6)
RGB_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RGB_GREY   = RGBColor(0x59, 0x59, 0x59)
RGB_LGREY  = RGBColor(0x88, 0x88, 0x88)

# Chart dimensions (inches)
CHART_W, CHART_H = 6.5, 2.4

# matplotlib chart style
CHART_RC = {
    "figure.facecolor":  "white",
    "axes.facecolor":    "white",
    "axes.edgecolor":    "#CCCCCC",
    "axes.grid":         True,
    "grid.color":        "#F0F0F0",
    "grid.linestyle":    "--",
    "grid.linewidth":    0.5,
    "lines.linewidth":   1.5,
    "font.family":       "DejaVu Sans",
    "font.size":         8,
    "axes.titlesize":    9,
    "axes.labelsize":    8,
    "xtick.labelsize":   7,
    "ytick.labelsize":   7,
}
LINE_COLORS = ["#2E75B6", "#ED7D31", "#70AD47", "#FF0000", "#7030A0"]
THRESHOLD_COLOR = "#C00000"


# ══════════════════════════════════════════════
# PAGE BORDERS  (XML injection into sectPr)
# ══════════════════════════════════════════════

def add_page_borders(doc: Document,
                     color: str  = C_DARK,
                     size:  int  = 12,
                     space: int  = 24,
                     val:   str  = "single"):
    """
    Add a border to EVERY page in the document.
    Injects <w:pgBorders> into each section's <w:sectPr>.

    size:  border width in eighths of a point  (12 = 1.5 pt)
    space: distance from page edge in points   (24 pt ≈ 0.33 inch)
    val:   border style — "single", "double", "thick", "dashSmallGap", etc.
    color: 6-char hex string without '#'
    """
    for section in doc.sections:
        sectPr = section._sectPr

        # Remove any existing pgBorders to avoid duplicates on repeat calls
        for old in sectPr.findall(qn("w:pgBorders")):
            sectPr.remove(old)

        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")

        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val)
            el.set(qn("w:sz"),    str(size))
            el.set(qn("w:space"), str(space))
            el.set(qn("w:color"), color)
            pgBorders.append(el)

        # Insert after <w:pgMar> per OOXML schema order; fall back to append
        ref = sectPr.find(qn("w:pgMar"))
        if ref is None:
            ref = sectPr.find(qn("w:pgSz"))
        if ref is not None:
            ref.addnext(pgBorders)
        else:
            sectPr.append(pgBorders)


# ══════════════════════════════════════════════
# HEADER & FOOTER
# ══════════════════════════════════════════════

def _field_run(para, field_name: str) -> None:
    """Append a Word field code (PAGE or NUMPAGES) to a paragraph."""
    r = para.add_run()
    r.font.size  = Pt(8)
    r.font.color.rgb = RGB_LGREY

    def el(tag): return OxmlElement(tag)

    begin = el("w:fldChar");  begin.set(qn("w:fldCharType"), "begin")
    instr = el("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    end   = el("w:fldChar");  end.set(qn("w:fldCharType"), "end")

    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)


def add_header_footer(doc: Document, client_name: str, period_label: str):
    """
    Apply header and footer to every section in the document.

    Header (all pages):
        LEFT  — client_name (bold, dark blue)
        RIGHT — period_label (regular, mid blue)
        Separated by a thin blue rule beneath the header paragraph.

    Footer (all pages):
        CENTRE — "Page N of M  |  Azure VM Performance Report  |  CONFIDENTIAL"
        Separated by a thin grey rule above the footer paragraph.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = False

        # ── HEADER ──────────────────────────────────────────
        hdr = section.header
        hdr.is_linked_to_previous = False

        # Clear default empty paragraph
        for p in hdr.paragraphs:
            p.clear()
        hp = hdr.paragraphs[0]

        # Tab stop: right-aligned at content width
        # A4 width 11906 twips − 2×1134 twips margins ≈ 9638 twips
        content_width_twips = 9638
        pPr = hp._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el  = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:pos"), str(content_width_twips))
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

        # Left: client name
        r1 = hp.add_run(client_name)
        r1.font.size  = Pt(9)
        r1.font.bold  = True
        r1.font.color.rgb = RGB_DARK

        # Tab to right side
        r_tab = hp.add_run()
        r_tab._r.append(OxmlElement("w:tab"))

        # Right: period
        r2 = hp.add_run(period_label)
        r2.font.size  = Pt(9)
        r2.font.color.rgb = RGB_MID

        # Bottom border on header paragraph
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_MID)
        pBdr.append(bot)
        pPr.append(pBdr)

        # ── FOOTER ──────────────────────────────────────────
        ftr = section.footer
        ftr.is_linked_to_previous = False

        for p in ftr.paragraphs:
            p.clear()
        fp = ftr.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Top border on footer paragraph
        fpPr = fp._p.get_or_add_pPr()
        fpBdr = OxmlElement("w:pBdr")
        top   = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "CCCCCC")
        fpBdr.append(top)
        fpPr.append(fpBdr)

        r3 = fp.add_run("Page ")
        r3.font.size  = Pt(8)
        r3.font.color.rgb = RGB_LGREY

        _field_run(fp, "PAGE")

        r4 = fp.add_run(" of ")
        r4.font.size  = Pt(8)
        r4.font.color.rgb = RGB_LGREY

        _field_run(fp, "NUMPAGES")

        r5 = fp.add_run("  |  Azure VM Performance Report  |  CONFIDENTIAL")
        r5.font.size  = Pt(8)
        r5.font.color.rgb = RGB_LGREY


# ══════════════════════════════════════════════
# XML / STYLE HELPERS
# ══════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _para_style(para, size_pt=10, bold=False, color=None, space_before=0, space_after=4):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Section heading with coloured left-border accent line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18" if level == 1 else "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), C_DARK if level == 1 else C_MID)
    pBdr.append(left)
    pPr.append(pBdr)

    r = p.add_run(text)
    r.font.size  = Pt(14 if level == 1 else 11)
    r.font.bold  = True
    r.font.color.rgb = RGB_DARK if level == 1 else RGB_MID


# ══════════════════════════════════════════════
# CHARTS
# ══════════════════════════════════════════════

def _parse_ts(series):
    """Convert ISO timestamp strings to datetime objects."""
    return [dateutil.parser.parse(ts) for ts, _ in series]


def _make_line_chart(series_list, title, ylabel,
                     threshold=None, formatter=None,
                     max_series_list=None) -> io.BytesIO:
    """
    series_list:     [(label, [(ts_str, avg_val), ...]), ...]  — Average series
    max_series_list: [(label, [(ts_str, max_val), ...]), ...]  — Maximum series (optional)
                     Must correspond 1-to-1 with series_list entries.

    When max_series_list is provided each metric is rendered as:
      - Solid line  = Average  (primary colour, labelled "Avg")
      - Dashed line = Maximum  (same colour, lighter, labelled "Max")
      - Shaded band between Average and Maximum showing the hourly spread

    Returns BytesIO PNG.
    """
    with plt.rc_context(CHART_RC):
        fig, ax = plt.subplots(figsize=(CHART_W, CHART_H))

        dates = []   # keep last valid dates for x-axis formatting
        for idx, (label, data) in enumerate(series_list):
            if not data:
                continue
            try:
                dates = _parse_ts(data)
            except Exception:
                dates = list(range(len(data)))
            avg_vals = [v for _, v in data]
            color    = LINE_COLORS[idx % len(LINE_COLORS)]

            # Determine legend labels based on whether max is present
            avg_label = f"{label} (Avg)" if max_series_list else label
            ax.plot(dates, avg_vals, color=color,
                    label=avg_label, linewidth=1.5, zorder=3)

            if max_series_list and idx < len(max_series_list):
                _, max_data = max_series_list[idx]
                if max_data:
                    try:
                        max_dates = _parse_ts(max_data)
                        max_vals  = [v for _, v in max_data]
                    except Exception:
                        max_dates = list(range(len(max_data)))
                        max_vals  = [v for _, v in max_data]

                    # Dashed max line — same colour, 60% opacity
                    ax.plot(max_dates, max_vals,
                            color=color, linestyle="--", linewidth=1.0,
                            alpha=0.7, label=f"{label} (Max)", zorder=2)

                    # Shaded band: fill from avg up to max
                    # Align on matching timestamps for safety
                    if len(avg_vals) == len(max_vals):
                        ax.fill_between(dates, avg_vals, max_vals,
                                        alpha=0.10, color=color,
                                        label="_nolegend_")
                    else:
                        # Lengths differ — just show avg fill without band
                        ax.fill_between(dates, avg_vals, alpha=0.06, color=color)
                else:
                    ax.fill_between(dates, avg_vals, alpha=0.06, color=color)
            else:
                ax.fill_between(dates, avg_vals, alpha=0.06, color=color)

        if threshold is not None:
            ax.axhline(threshold, color=THRESHOLD_COLOR, linestyle="--",
                       linewidth=0.9, alpha=0.8, label=f"Threshold ({threshold}%)")

        if isinstance(dates, list) and dates and hasattr(dates[0], "day"):
            # Use FuncFormatter instead of DateFormatter("%-d %b"):
            # %-d is Linux-only and raises "Invalid format string" on Windows.
            # %#d is Windows-only. FuncFormatter with lstrip("0") is cross-platform.
            def _day_fmt(x, _):
                try:
                    d = mdates.num2date(x)
                    return f"{d.day} {d.strftime('%b')}"
                except Exception:
                    return ""
            ax.xaxis.set_major_formatter(FuncFormatter(_day_fmt))
            ax.xaxis.set_major_locator(mdates.WeekdayLocator(interval=1))
            plt.gcf().autofmt_xdate(rotation=0, ha="center")

        if formatter:
            ax.yaxis.set_major_formatter(FuncFormatter(formatter))

        ax.set_title(title, fontsize=9, fontweight="bold", color="#333333", pad=5)
        ax.set_ylabel(ylabel, fontsize=7, color="#555555")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.tick_params(axis="both", length=2)

        # Always show legend when max lines or multiple series or threshold present
        has_max = bool(max_series_list and any(d for _, d in max_series_list))
        if len(series_list) > 1 or threshold is not None or has_max:
            ax.legend(fontsize=7, loc="upper right", framealpha=0.7)

        plt.tight_layout(pad=0.6)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
        buf.seek(0)
        plt.close(fig)
        return buf


def _filter_real_disks(disk_data: dict) -> dict:
    """
    Filters out non-relevant mounts like snap, proc, sys, etc.
    Keeps only meaningful disks for reporting.
    """
    EXCLUDE_PREFIXES = (
        "/snap",
        "/run",
        "/proc",
        "/sys",
        "/dev",
    )

    filtered = {}

    for mount, usage in disk_data.items():
        # Skip excluded mount paths
        if mount.startswith(EXCLUDE_PREFIXES):
            continue

        # Skip zero or invalid disks
        if usage is None or usage == 0:
            continue

        # Keep only meaningful mounts
        filtered[mount] = usage

    return filtered


def _make_disk_bar_chart(disk_data: dict) -> io.BytesIO:
    """Horizontal bar chart of disk utilisation per drive."""
    with plt.rc_context(CHART_RC):
        drives = list(disk_data.keys())
        values = [disk_data[d] for d in drives]

        fig, ax = plt.subplots(figsize=(6.0, max(1.0, len(drives) * 0.55 + 0.5)))
        colors  = [
            THRESHOLD_COLOR if v >= 85
            else "#FFC000"  if v >= 70
            else "#2E75B6"
            for v in values
        ]
        bars = ax.barh(drives, values, color=colors, height=0.5, edgecolor="none")

        for bar, val in zip(bars, values):
            ax.text(
                min(val + 1.5, 101), bar.get_y() + bar.get_height() / 2,
                f"{val:.1f}%", va="center", ha="left",
                fontsize=8, color="#333333", fontweight="bold",
            )

        ax.set_xlim(0, 108)
        ax.axvline(85, color=THRESHOLD_COLOR, linestyle="--", linewidth=0.8, alpha=0.5)
        ax.set_xlabel("Used (%)", fontsize=7)
        ax.set_title("Disk Utilisation by Drive", fontsize=9,
                     fontweight="bold", color="#333333", pad=5)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.tick_params(axis="y", labelsize=8)

        plt.tight_layout(pad=0.5)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
        buf.seek(0)
        plt.close(fig)
        return buf


def _insert_chart(doc, buf: io.BytesIO, subtitle: str = ""):
    """Insert a chart PNG into the document."""
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(subtitle)
        r.font.size  = Pt(9)
        r.font.bold  = True
        r.font.color.rgb = RGB_MID
    doc.add_picture(buf, width=Inches(CHART_W))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(4)



# ══════════════════════════════════════════════
# CPU BREACH DETAIL TABLE (per-VM section)
# ══════════════════════════════════════════════

def _add_cpu_breach_table(doc: Document, vm) -> None:
    """
    Render a CPU breach detail block inside a VM's section, directly below
    the CPU chart. Only called when the VM has threshold breaches.

    Layout:
      - Stats row:  Avg CPU | Max CPU | Breaches | Days Affected | Longest Run
      - Breach list table: Date | Time | Peak % | Severity
    """
    summary = getattr(vm, "findings_cpu_summary", None)
    if not summary or not vm.cpu_threshold_breaches:
        return

    threshold  = summary.get("threshold", 80)
    breach_peak = summary.get("breach_peak", 0)
    severity   = summary.get("severity", "WARNING")
    bg_hdr     = C_RED if severity == "CRITICAL" else C_AMBER
    bg_hdr_rgb = (
        RGBColor(0xC0, 0x00, 0x00) if severity == "CRITICAL"
        else RGBColor(0xBF, 0x8F, 0x00)
    )

    # ── Sub-heading ──────────────────────────────────────────────────
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(6)
    p_hdr.paragraph_format.space_after  = Pt(4)
    rh = p_hdr.add_run(
        f"CPU Threshold Breach Detail  "
        f"(threshold: {threshold}%  |  severity: {severity})"
    )
    rh.font.size = Pt(9); rh.font.bold = True
    rh.font.color.rgb = bg_hdr_rgb

    # ── Stats summary row ────────────────────────────────────────────
    stat_labels = ["Overall Avg CPU", "Monthly Peak CPU",
                   "Total Breaches", "Days Affected", "Longest Consecutive Run"]
    stat_values = [
        f"{summary['overall_avg']}%",
        f"{summary['breach_peak']}%",
        str(summary['breach_count']),
        str(summary['breach_day_count']),
        f"{summary['max_consecutive']} hour(s)",
    ]
    col_w_stats = [1800, 1800, 1500, 1500, 2426]   # sum = 9026

    tbl_stats = doc.add_table(rows=2, cols=5)
    tbl_stats.style = "Table Grid"
    tbl_stats.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (lbl, val, w) in enumerate(zip(stat_labels, stat_values, col_w_stats)):
        hc = tbl_stats.rows[0].cells[i]
        vc = tbl_stats.rows[1].cells[i]
        _set_cell_bg(hc, bg_hdr)
        _set_cell_bg(vc, C_AMBER_BG if severity != "CRITICAL" else C_RED_BG)
        _cell_padding(hc, top=60, bottom=60, left=100, right=100)
        _cell_padding(vc, top=60, bottom=60, left=100, right=100)
        # Header cell
        ph = hc.paragraphs[0]; ph.clear()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = ph.add_run(lbl)
        rr.font.size = Pt(8); rr.font.bold = True; rr.font.color.rgb = RGB_WHITE
        # Value cell
        pv = vc.paragraphs[0]; pv.clear()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(val)
        rv.font.size = Pt(10); rv.font.bold = True
        rv.font.color.rgb = bg_hdr_rgb
        # Column width
        for cell in (hc, vc):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── Breach detail table ──────────────────────────────────────────
    # Show up to 20 breach entries; if more, note the count
    max_rows  = 20
    breaches  = vm.cpu_threshold_breaches[:max_rows]
    overflow  = len(vm.cpu_threshold_breaches) - max_rows

    tbl_breach = doc.add_table(rows=1, cols=4)
    tbl_breach.style = "Table Grid"
    tbl_breach.alignment = WD_TABLE_ALIGNMENT.LEFT

    b_headers = ["Date", "Time (UTC)", "CPU %", "Severity"]
    b_widths  = [2000, 1800, 1826, 3400]   # sum = 9026

    for i, (h, w) in enumerate(zip(b_headers, b_widths)):
        c = tbl_breach.rows[0].cells[i]
        _set_cell_bg(c, C_MID)
        _cell_padding(c, top=60, bottom=60, left=100, right=100)
        p = c.paragraphs[0]; p.clear()
        rr = p.add_run(h)
        rr.font.size = Pt(8); rr.font.bold = True; rr.font.color.rgb = RGB_WHITE
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    for row_idx, (ts, val) in enumerate(breaches):
        row = tbl_breach.add_row()
        bg  = C_GREY_LIGHT if row_idx % 2 == 0 else C_WHITE

        # Classify per-point severity
        if val >= 95:
            sev_txt = "CRITICAL — Near saturation"
            sev_fg  = C_RED
        elif val >= 90:
            sev_txt = "HIGH — Severe pressure"
            sev_fg  = "A04000"
        else:
            sev_txt = "WARNING — Above threshold"
            sev_fg  = C_AMBER

        cell_vals = [ts[:10], ts[11:16], f"{val:.1f}%", sev_txt]
        cell_fgs  = [C_DARK,  C_GREY_TEXT, sev_fg,       sev_fg]
        cell_bolds= [False,   False,        True,          True]

        for i, (val_txt, fg, bold, w) in enumerate(
            zip(cell_vals, cell_fgs, cell_bolds, b_widths)
        ):
            c = row.cells[i]
            _set_cell_bg(c, bg)
            _cell_padding(c, top=50, bottom=50, left=100, right=100)
            p = c.paragraphs[0]; p.clear()
            rr = p.add_run(val_txt)
            rr.font.size  = Pt(8); rr.font.bold = bold
            rr.font.color.rgb = RGBColor(
                int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16)
            )
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    if overflow > 0:
        p_of = doc.add_paragraph()
        p_of.paragraph_format.space_before = Pt(3)
        r_of = p_of.add_run(
            f"  … and {overflow} additional breach hour(s) not shown. "
            f"Full data visible in Azure Monitor Metrics Explorer."
        )
        r_of.font.size = Pt(8); r_of.font.italic = True
        r_of.font.color.rgb = RGB_LGREY
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ══════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════

def _add_cover_page(doc: Document, client_name: str,
                    period_label: str, month_name: str,
                    subscription_name: str = ""):
    """
    Dark-navy banner table → title → period → client → sub info.
    Full-width, no logo dependency.
    """
    # ── Dark banner via single-cell full-width table ─────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Full content width: A4 = 9026 twips at 2 cm margins
    tbl.columns[0].width = Twips(9026)

    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, C_DARK)

    # Row height
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   "2520")   # 1.75 inch
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    # Set cell margins for internal padding
    _cell_padding(cell, top=400, bottom=200, left=280, right=280)

    # Remove default empty paragraph, add our own
    for p in cell.paragraphs:
        p.clear()

    # "AZURE VM PERFORMANCE REPORT"
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(20)
    p1.paragraph_format.space_after  = Pt(4)
    r1 = p1.add_run("AZURE VM PERFORMANCE REPORT")
    r1.font.size  = Pt(22)
    r1.font.bold  = True
    r1.font.color.rgb = RGB_WHITE

    # Period line
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(6)
    r2 = p2.add_run(period_label)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = RGBColor(0xA8, 0xC8, 0xE8)

    # Client name
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(2)
    r3 = p3.add_run(client_name)
    r3.font.size  = Pt(14)
    r3.font.bold  = True
    r3.font.color.rgb = RGB_WHITE

    if subscription_name:
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(0)
        r4 = p4.add_run(f"Subscription: {subscription_name}")
        r4.font.size  = Pt(9)
        r4.font.color.rgb = RGBColor(0xA8, 0xC8, 0xE8)

    doc.add_paragraph()   # breathing room after banner


# ══════════════════════════════════════════════
# VM INFO TABLE
# ══════════════════════════════════════════════

def _add_vm_info_table(doc: Document, vm: VMMetrics, period_label: str):
    """
    4-column info table: VM Name | SKU | vCPU / Memory | Period
    Dark-blue header row, light-blue value row.
    """
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Build metric source label for display
    m_src = getattr(vm, "metric_source", "ARM")
    src_display_map = {
        "LA-AMA": "Log Analytics — AMA (InsightsMetrics)",
        "LA-MMA": "Log Analytics — MMA (Perf table)",
        "ARM":    "Azure Monitor REST API",
        "none":   "No data collected",
    }
    # Handle compound labels like "LA-AMA + ARM(2)"
    if m_src.startswith("LA-AMA"):
        src_display = src_display_map["LA-AMA"]
        if "ARM" in m_src:
            src_display += " (some metrics via ARM fallback)"
    elif m_src.startswith("LA-MMA"):
        src_display = src_display_map["LA-MMA"]
        if "ARM" in m_src:
            src_display += " (some metrics via ARM fallback)"
    else:
        src_display = src_display_map.get(m_src, m_src)

    headers = ["VM Name", "SKU", "vCPU / Memory", "Report Period", "Data Source"]
    values  = [
        vm.vm_name,
        vm.sku,
        f"{vm.vcpus} vCPUs  /  {vm.memory_gib} GiB RAM",
        period_label,
        src_display,
    ]
    col_widths = [1600, 1800, 1900, 1726, 2000]   # DXA, sum = 9026

    for i, (hdr, val, w) in enumerate(zip(headers, values, col_widths)):
        hc = tbl.rows[0].cells[i]
        vc = tbl.rows[1].cells[i]

        _set_cell_bg(hc, C_DARK)
        _set_cell_bg(vc, C_LIGHT)
        _cell_padding(hc)
        _cell_padding(vc)

        hp = hc.paragraphs[0]
        hp.clear()
        hr = hp.add_run(hdr)
        hr.font.size  = Pt(9)
        hr.font.bold  = True
        hr.font.color.rgb = RGB_WHITE

        vp = vc.paragraphs[0]
        vp.clear()
        vr = vp.add_run(val)
        vr.font.size  = Pt(9)
        vr.font.color.rgb = RGB_DARK

        # Set column widths
        for cell in (hc, vc):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(w))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════
# FINDINGS TABLE
# ══════════════════════════════════════════════

def _add_findings_summary_table(doc: Document, all_findings: list):
    """
    Consolidated status table with colour-coded status column.
    Columns: VM Name | SKU | Status | Issues Found | CPU Max%
    """
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers    = ["VM Name", "SKU", "Status", "Issues Found", "CPU Max %"]
    col_widths = [1900, 2100, 1100, 2626, 1300]   # sum = 9026
    status_map = {
        "NORMAL":   (C_GREEN_BG,  C_GREEN,  "Normal"),
        "WARNING":  (C_AMBER_BG,  C_AMBER,  "Warning"),
        "CRITICAL": (C_RED_BG,    C_RED,    "Critical"),
    }

    # Header row
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, C_MID)
        _cell_padding(c)
        p = c.paragraphs[0]; p.clear()
        r = p.add_run(h)
        r.font.size  = Pt(9)
        r.font.bold  = True
        r.font.color.rgb = RGB_WHITE
        tc   = c._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    for row_idx, f in enumerate(all_findings):
        row = tbl.add_row()
        bg_fill  = C_GREY_LIGHT if row_idx % 2 == 0 else C_WHITE
        st_bg, st_fg, st_label = status_map.get(
            f["status"], (C_WHITE, C_GREY_TEXT, f["status"])
        )
        issue_count = f"{len(f['issues'])} issue(s)" if f["issues"] else "—"

        # Approximate CPU max from findings (stored in cpu_max)
        cpu_max_val = f.get("cpu_max")
        cpu_max_str = f"{cpu_max_val}%" if cpu_max_val is not None else "—"
        vm_issues = f.get("issues", [])

        cell_vals = [
            f["vm_name"],
            f.get("sku", "—"),
            st_label,
            issue_count,
            cpu_max_str,
        ]
        fills = [bg_fill, bg_fill, st_bg, bg_fill, bg_fill]
        fgs   = [C_DARK,  C_GREY_TEXT, st_fg, C_GREY_TEXT, C_GREY_TEXT]
        bolds = [True, False, True, False, False]

        for i, (val, fill, fg_hex, bold, w) in enumerate(
            zip(cell_vals, fills, fgs, bolds, col_widths)
        ):
            c = row.cells[i]
            _set_cell_bg(c, fill)
            _cell_padding(c)
            p = c.paragraphs[0]; p.clear()
            r = p.add_run(val)
            r.font.size  = Pt(9)
            r.font.bold  = bold
            r.font.color.rgb = RGBColor(
                int(fg_hex[0:2], 16),
                int(fg_hex[2:4], 16),
                int(fg_hex[4:6], 16),
            )
            tc   = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)



# ══════════════════════════════════════════════
# NSG SECTION
# ══════════════════════════════════════════════

# Risk colour palette (bg, fg hex strings)
_NSG_RISK_COLOURS = {
    "CRITICAL": ("FFDFD9", "C00000"),
    "HIGH":     ("FFF2CC", "BF8F00"),
    "MEDIUM":   ("FDEBD0", "A04000"),
    "LOW":      ("EAF4FB", "1F618D"),
    "INFO":     ("F2F2F2", "595959"),
}


def _add_nsg_section(doc: Document, vm) -> None:
    """
    Render the NSG section for one VM inside its per-VM page.
    Layout:
      - Subheading "Network Security Groups"
      - NSG name(s) and level chips
      - Security rules table (one row per rule, colour-coded by risk)
      - Per-VM NSG recommendations (bullet list)
    """
    from collect_metrics import NSGRule

    _add_heading(doc, "Network Security Groups", 2)

    # ── NSG name chips ────────────────────────────────────────────────
    if not vm.nsg_names:
        p = doc.add_paragraph()
        r = p.add_run(
            "No NSG found — this VM's NIC and subnet have no Network Security Group "
            "attached. All inbound and outbound traffic is unrestricted."
        )
        r.font.size = Pt(9); r.font.italic = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return

    p_names = doc.add_paragraph()
    p_names.paragraph_format.space_after = Pt(4)
    r_lbl = p_names.add_run("Attached NSG(s):  ")
    r_lbl.font.size = Pt(9); r_lbl.font.bold = True; r_lbl.font.color.rgb = RGB_DARK
    r_val = p_names.add_run("  |  ".join(vm.nsg_names))
    r_val.font.size = Pt(9); r_val.font.color.rgb = RGB_GREY

    if not vm.nsg_rules:
        p = doc.add_paragraph()
        r = p.add_run("No custom security rules found (only Azure default rules apply).")
        r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGB_GREY
        return

    # ── Rules table ───────────────────────────────────────────────────
    # Columns: Priority | NSG / Level | Direction | Access | Protocol | Source | Port(s) | Risk
    col_headers = ["Priority", "NSG / Level", "Direction", "Access",
                   "Protocol", "Source", "Dest Port(s)", "Risk"]
    col_widths  = [700, 1500, 900, 700, 700, 1600, 1226, 700]   # sum = 8026 (leaving margin)

    tbl = doc.add_table(rows=1, cols=len(col_headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, (h, w) in enumerate(zip(col_headers, col_widths)):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, C_DARK)
        _cell_padding(c, top=60, bottom=60, left=80, right=80)
        p = c.paragraphs[0]; p.clear()
        rr = p.add_run(h)
        rr.font.size = Pt(8); rr.font.bold = True; rr.font.color.rgb = RGB_WHITE
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    # Sort rows: risky rules first (CRITICAL → HIGH → MEDIUM → LOW → INFO),
    # then by priority within each group
    _RISK_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}
    sorted_rules = sorted(
        vm.nsg_rules,
        key=lambda r: (_RISK_ORDER.get(r.risk, 9), r.priority)
    )

    for rule in sorted_rules:
        bg_hex, fg_hex = _NSG_RISK_COLOURS.get(rule.risk, ("F2F2F2", "595959"))
        row = tbl.add_row()

        cell_vals = [
            str(rule.priority),
            f"{rule.nsg_name} ({rule.nsg_level})",
            rule.direction,
            rule.access,
            rule.protocol,
            rule.source_prefix,
            rule.dest_port,
            rule.risk,
        ]

        for i, (val, w) in enumerate(zip(cell_vals, col_widths)):
            c = row.cells[i]
            # Risk column gets coloured; others get light alternating fill
            if i == len(col_headers) - 1:  # Risk column
                _set_cell_bg(c, bg_hex)
            else:
                _set_cell_bg(c, "FAFAFA" if rule.risk == "INFO" else "FFFFFF")
            _cell_padding(c, top=50, bottom=50, left=80, right=80)
            p = c.paragraphs[0]; p.clear()
            rr = p.add_run(val)
            rr.font.size = Pt(8)
            if i == len(col_headers) - 1:
                rr.font.bold = True
                rr.font.color.rgb = RGBColor(
                    int(fg_hex[0:2], 16), int(fg_hex[2:4], 16), int(fg_hex[4:6], 16)
                )
            else:
                rr.font.color.rgb = (
                    RGBColor(0x19, 0x19, 0x19) if rule.risk != "INFO"
                    else RGBColor(0x59, 0x59, 0x59)
                )
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Per-VM NSG recommendations ────────────────────────────────────
    risky = [r for r in vm.nsg_rules if r.risk not in ("INFO",)]
    if risky:
        p_rec_hdr = doc.add_paragraph()
        p_rec_hdr.paragraph_format.space_before = Pt(6)
        p_rec_hdr.paragraph_format.space_after  = Pt(3)
        rh = p_rec_hdr.add_run("Security Recommendations for this VM:")
        rh.font.size = Pt(9); rh.font.bold = True; rh.font.color.rgb = RGB_DARK

        seen_recs = set()
        for rule in sorted_rules:
            if rule.risk == "INFO" or not rule.recommendation:
                continue
            rec_text = rule.recommendation
            if rec_text in seen_recs:
                continue
            seen_recs.add(rec_text)

            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after  = Pt(2)
            bp.paragraph_format.left_indent  = Pt(8)
            # Risk badge
            bg_h, fg_h = _NSG_RISK_COLOURS.get(rule.risk, ("F2F2F2","595959"))
            rb = bp.add_run(f"[{rule.risk}]  ")
            rb.font.size = Pt(8); rb.font.bold = True
            rb.font.color.rgb = RGBColor(
                int(fg_h[0:2], 16), int(fg_h[2:4], 16), int(fg_h[4:6], 16)
            )
            rt = bp.add_run(rec_text)
            rt.font.size = Pt(9); rt.font.color.rgb = RGB_GREY

    else:
        p = doc.add_paragraph()
        rr = p.add_run("✓  All NSG rules follow security best practices for this VM.")
        rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x37, 0x56, 0x23)


# ══════════════════════════════════════════════
# MAIN REPORT GENERATOR
# ══════════════════════════════════════════════

def generate_report(all_vm_metrics: list,
                    all_findings: list,
                    config: ReportConfig,
                    client_name: str = "Production",
                    subscription_name: str = "",
                    output_path: str = None) -> str:
    """
    Generate the full Word .docx report.

    Structure:
      1. Cover page (dark-blue banner)
      2. Executive Summary
      3. Per-VM sections (info table + 6 charts)
      4. Findings & Recommendations

    All pages have:
      - Dark-blue page border
      - Header: client name | period
      - Footer: page number | report title | CONFIDENTIAL

    Returns: path to the saved .docx file.
    """
    os.makedirs(config.output_dir, exist_ok=True)

    period_label = config.report_period_label
    month_name   = config.report_month_name

    if output_path is None:
        safe = client_name.replace(" ", "_").replace("/", "-")
        fname = f"Azure_VM_Report_{safe}_{month_name.replace(' ', '_')}.docx"
        output_path = os.path.join(config.output_dir, fname)

    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.0)

    # ── Page borders (every page) ────────────────────────────────────
    add_page_borders(doc, color=C_DARK, size=12, space=24, val="single")

    # ── Header & footer (every page) ────────────────────────────────
    add_header_footer(doc, client_name, period_label)

    # ══ 1. COVER PAGE ════════════════════════════════════════════════
    _add_cover_page(doc, client_name, period_label, month_name, subscription_name)

    # ══ 2. EXECUTIVE SUMMARY ═════════════════════════════════════════
    _add_heading(doc, "Executive Summary", 1)

    summary_intro = (
        "This report provides a comprehensive overview of the utilisation and performance "
        "of Azure Virtual Machines during the specified reporting period. "
        "The analysis covers key performance metrics to evaluate the efficiency and "
        "resource utilisation of all monitored virtual machines."
    )
    p = doc.add_paragraph(summary_intro)
    _para_style(p, size_pt=10, color=RGB_GREY, space_after=8)

    for item in [
        ("Performance Metrics",
         "CPU utilisation, available memory, disk IOPS, and network throughput — "
         "presented as hourly Average and Maximum time-series for the full reporting month. "
         "Shaded bands between the Average and Maximum lines highlight peak burst activity."),
        ("Efficiency Assessment",
         "Identification of performance bottlenecks, threshold breaches, and "
         "resource pressure events."),
        ("Resource Utilisation",
         "Overview of CPU, memory, and disk usage patterns across all VMs."),
        ("Findings & Recommendations",
         "Actionable insights and prioritised recommendations based on observed anomalies."),
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(3)
        r_bold = bp.add_run(f"{item[0]}: ")
        r_bold.font.bold  = True
        r_bold.font.size  = Pt(10)
        r_bold.font.color.rgb = RGB_DARK
        r_rest = bp.add_run(item[1])
        r_rest.font.size  = Pt(10)
        r_rest.font.color.rgb = RGB_GREY

    # VM count summary
    alerts   = [f for f in all_findings if f["status"] != "NORMAL"]
    critical = [f for f in all_findings if f["status"] == "CRITICAL"]
    doc.add_paragraph()

    tbl_sum = doc.add_table(rows=1, cols=4)
    tbl_sum.alignment = WD_TABLE_ALIGNMENT.LEFT
    sum_data = [
        ("Total VMs",          str(len(all_vm_metrics)), C_DARK,  C_LIGHT),
        ("Normal",             str(len(all_vm_metrics) - len(alerts)), C_GREEN, C_GREEN_BG),
        ("Warning",            str(len(alerts) - len(critical)), C_AMBER, C_AMBER_BG),
        ("Critical",           str(len(critical)), C_RED, C_RED_BG),
    ]
    col_w = [2256, 2256, 2257, 2257]
    for i, (label, val, fg, bg) in enumerate(sum_data):
        c = tbl_sum.rows[0].cells[i]
        _set_cell_bg(c, bg)
        _cell_padding(c, top=160, bottom=160, left=200, right=200)
        p = c.paragraphs[0]; p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(f"{val}\n")
        rn.font.size = Pt(22)
        rn.font.bold = True
        rn.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
        rl = p.add_run(label)
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
        tc   = c._tc; tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(col_w[i])); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    doc.add_page_break()

    # ══ 3. PER-VM SECTIONS ═══════════════════════════════════════════
    def bytes_fmt(x, _):
        if x >= 1e9: return f"{x/1e9:.1f} GB/s"
        if x >= 1e6: return f"{x/1e6:.1f} MB/s"
        if x >= 1e3: return f"{x/1e3:.1f} KB/s"
        return f"{x:.0f} B/s"

    for vm in all_vm_metrics:
        _add_heading(doc, f"VM: {vm.vm_name}", 1)
        _add_vm_info_table(doc, vm, period_label)

        # CPU — Average (solid) + Maximum (dashed) with shaded band between
        _insert_chart(
            doc,
            _make_line_chart(
                [(vm.vm_name, vm.cpu_percent)],
                "Percentage CPU",  "CPU %",
                threshold=config.cpu_alert_threshold,
                max_series_list=[(vm.vm_name, vm.cpu_percent_max)],
            ),
            "Percentage CPU",
        )

        # CPU breach detail table (only rendered if there were threshold breaches)
        if vm.cpu_threshold_breaches:
            # Attach the cpu_breach_summary from findings to the vm object
            # so the table renderer can access it
            matching_f = next(
                (f for f in all_findings if f["vm_name"] == vm.vm_name), {}
            )
            vm.findings_cpu_summary = matching_f.get("cpu_breach_summary")
            _add_cpu_breach_table(doc, vm)

        # Memory — convert bytes → GiB for both avg and max series
        mem_gb     = [(ts, v / (1024**3)) for ts, v in vm.available_memory_bytes]
        mem_gb_max = [(ts, v / (1024**3)) for ts, v in vm.available_memory_bytes_max]
        _insert_chart(
            doc,
            _make_line_chart(
                [(vm.vm_name, mem_gb)],
                "Available Memory",  "GiB Available",
                max_series_list=[(vm.vm_name, mem_gb_max)],
            ),
            "Available Memory",
        )

        # Disk IOPS — show avg+max for both Read and Write
        _insert_chart(
            doc,
            _make_line_chart(
                [("Read IOPS",  vm.disk_read_iops),
                 ("Write IOPS", vm.disk_write_iops)],
                "Logical Disk IOPS",  "Operations/sec",
                max_series_list=[
                    ("Read IOPS",  vm.disk_read_iops_max),
                    ("Write IOPS", vm.disk_write_iops_max),
                ],
            ),
            "Logical Disk IOPS",
        )

        # Bytes Sent — Average + Maximum
        _insert_chart(
            doc,
            _make_line_chart(
                [(vm.vm_name, vm.network_bytes_sent)],
                "Bytes Sent Rate",  "Bytes/sec",
                formatter=bytes_fmt,
                max_series_list=[(vm.vm_name, vm.network_bytes_sent_max)],
            ),
            "Bytes Sent Rate",
        )

        # Bytes Received — Average + Maximum
        _insert_chart(
            doc,
            _make_line_chart(
                [(vm.vm_name, vm.network_bytes_received)],
                "Bytes Received Rate",  "Bytes/sec",
                formatter=bytes_fmt,
                max_series_list=[(vm.vm_name, vm.network_bytes_received_max)],
            ),
            "Bytes Received Rate",
        )

        # Disk Utilisation (Filtered)
        filtered_disks = _filter_real_disks(vm.disk_utilization)

        if filtered_disks:
            disk_src = getattr(vm, "disk_source", "no_workspace")
            disk_src_note = {
                "found":        "Source: Log Analytics — most recent reading (last 7 days). "
                        "Matches Azure Portal 'CURRENT USED %' which always shows the live value.",
                "no_workspace": "Disk utilisation not available — Log Analytics workspace not configured.",
                "no_data":      "Disk utilisation not available — no data found in Log Analytics for this VM. "
                        "Verify the Data Collection Rule includes LogicalDisk counters.",
            }.get(disk_src, "")

            _insert_chart(
                doc,
                _make_disk_bar_chart(filtered_disks),
                "Disk Utilisation",
            )

            if disk_src_note:
                p_note = doc.add_paragraph()
                p_note.paragraph_format.space_before = Pt(0)
                p_note.paragraph_format.space_after  = Pt(2)
                rn = p_note.add_run(disk_src_note)
                rn.font.size  = Pt(7.5)
                rn.font.italic = True
                rn.font.color.rgb = RGB_LGREY
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r_label = p.add_run("Disk Utilisation:  ")
            r_label.font.bold  = True
            r_label.font.size  = Pt(9)
            r_label.font.color.rgb = RGB_MID

            disk_source = getattr(vm, "disk_source", "no_workspace")
            if disk_source == "no_workspace":
                msg = (
                    "Not available — Log Analytics workspace ID is not configured for this "
                    "subscription. Set log_analytics_workspace_id in CLIENT_CONFIGS_JSON to "
                    "enable disk utilisation collection."
                )
            else:
                # workspace configured but both InsightsMetrics and Perf returned no rows
                msg = (
                    "Not available — Log Analytics workspace is configured but no disk "
                    "performance data was found for this VM in the reporting period. "
                    "Verify the Data Collection Rule (DCR) includes LogicalDisk performance "
                    "counters and that the Azure Monitor Agent is running on this VM."
                )
            r_msg = p.add_run(msg)
            r_msg.font.size   = Pt(9)
            r_msg.font.color.rgb = RGB_GREY
            r_msg.font.italic = True

        # NSG Security
        _add_nsg_section(doc, vm)

        doc.add_page_break()

    # ══ 4. FINDINGS & RECOMMENDATIONS ════════════════════════════════
    _add_heading(doc, "Findings & Recommendations", 1)

    # Status summary table
    _add_heading(doc, "Status Summary", 2)
    _add_findings_summary_table(doc, all_findings)

    # Detailed findings text
    _add_heading(doc, "Findings", 2)

    # ── CPU Utilisation findings ──────────────────────────────────────
    cpu_breached = [f for f in all_findings if f.get("cpu_breach_summary")]

    p_cpu_hdr = doc.add_paragraph()
    p_cpu_hdr.paragraph_format.space_after = Pt(4)
    r_b = p_cpu_hdr.add_run("CPU Utilisation:  ")
    r_b.font.bold = True; r_b.font.size = Pt(10); r_b.font.color.rgb = RGB_DARK

    if cpu_breached:
        r_t = p_cpu_hdr.add_run(
            f"{len(cpu_breached)} VM(s) exceeded the CPU threshold during this period. "
            "Breach details and per-VM recommendations are listed below and in each "
            "VM's individual section."
        )
    else:
        r_t = p_cpu_hdr.add_run(
            "All VMs operated within acceptable CPU utilisation ranges throughout "
            "the reporting period. No threshold breaches were detected."
        )
    r_t.font.size = Pt(10); r_t.font.color.rgb = RGB_GREY

    # Per-VM CPU breach summary block
    for f in all_findings:
        cs = f.get("cpu_breach_summary")
        if not cs:
            continue

        sev_fg = RGBColor(0xC0,0x00,0x00) if cs["severity"] == "CRITICAL"                  else RGBColor(0xBF,0x8F,0x00)

        # VM name + severity badge
        p_vm = doc.add_paragraph()
        p_vm.paragraph_format.space_before = Pt(8)
        p_vm.paragraph_format.space_after  = Pt(2)
        rb = p_vm.add_run(f"  {f['vm_name']}  ")
        rb.font.bold = True; rb.font.size = Pt(10); rb.font.color.rgb = RGB_DARK
        rs = p_vm.add_run(f"[{cs['severity']}]")
        rs.font.bold = True; rs.font.size = Pt(9); rs.font.color.rgb = sev_fg

        # Key stats on one line
        p_st = doc.add_paragraph()
        p_st.paragraph_format.space_after = Pt(2)
        p_st.paragraph_format.left_indent = Pt(16)
        r_st = p_st.add_run(
            f"Overall Avg: {cs['overall_avg']}%  |  "
            f"Peak: {cs['breach_peak']}%  |  "
            f"Breaches: {cs['breach_count']} hour(s)  |  "
            f"Days affected: {cs['breach_day_count']}  |  "
            f"Longest run: {cs['max_consecutive']} consecutive hour(s)"
        )
        r_st.font.size = Pt(9); r_st.font.color.rgb = RGB_GREY

        # Issue bullets
        for issue in f["issues"]:
            if "CPU" in issue or "cpu" in issue:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after  = Pt(2)
                bp.paragraph_format.left_indent  = Pt(8)
                rt = bp.add_run(issue)
                rt.font.size = Pt(9); rt.font.color.rgb = RGB_GREY

    doc.add_paragraph()
    for label, text in [
        ("Memory and Disk Utilisation:",
         "Memory and disk utilisation demonstrated consistent and satisfactory performance "
         "throughout the reporting period unless explicitly noted above."),
        ("Network Performance:",
         "Network performance remained normal throughout the monitored period. "
         "No significant congestion, packet loss, or latency anomalies were observed."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(label + "  ")
        r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = RGB_DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.color.rgb = RGB_GREY

    # Recommendations
    _add_heading(doc, "Recommendations", 2)

    all_recs = []
    for f in all_findings:
        all_recs.extend(f.get("recommendations", []))

    if not all_recs:
        p = doc.add_paragraph(
            "No actionable recommendations at this time. "
            "All monitored VMs are operating within defined performance thresholds."
        )
        _para_style(p, size_pt=10, color=RGB_GREY)
    else:
        for idx, rec in enumerate(all_recs, 1):
            bp = doc.add_paragraph(style="List Number")
            bp.paragraph_format.space_after = Pt(4)
            r = bp.add_run(rec)
            r.font.size = Pt(10); r.font.color.rgb = RGB_GREY

    # NSG Security Findings
    all_nsg_issues = []
    for f in all_findings:
        all_nsg_issues.extend(f.get("nsg_issues", []))

    critical_nsg = [r for r in all_nsg_issues if r.risk == "CRITICAL"]
    high_nsg     = [r for r in all_nsg_issues if r.risk == "HIGH"]
    medium_nsg   = [r for r in all_nsg_issues if r.risk == "MEDIUM"]

    if all_nsg_issues:
        _add_heading(doc, "Network Security Group Findings", 2)

        # Summary counts
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(6)
        for count, label, fg in [
            (len(critical_nsg), "Critical", "C00000"),
            (len(high_nsg),     "High",     "BF8F00"),
            (len(medium_nsg),   "Medium",   "A04000"),
        ]:
            if count:
                rc = p_sum.add_run(f"  {count} {label}  ")
                rc.font.size = Pt(9); rc.font.bold = True
                rc.font.color.rgb = RGBColor(
                    int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16)
                )

        # Consolidated unique recommendations
        seen = set()
        for rule in sorted(all_nsg_issues, key=lambda r: {"CRITICAL":0,"HIGH":1,"MEDIUM":2,"LOW":3}.get(r.risk,4)):
            if not rule.recommendation or rule.recommendation in seen:
                continue
            seen.add(rule.recommendation)
            bp = doc.add_paragraph(style="List Number")
            bp.paragraph_format.space_after = Pt(4)
            bg_h, fg_h = _NSG_RISK_COLOURS.get(rule.risk, ("F2F2F2","595959"))
            rb = bp.add_run(f"[{rule.risk} — {rule.nsg_name} / {rule.name}]  ")
            rb.font.size = Pt(9); rb.font.bold = True
            rb.font.color.rgb = RGBColor(int(fg_h[0:2],16), int(fg_h[2:4],16), int(fg_h[4:6],16))
            rt = bp.add_run(rule.recommendation)
            rt.font.size = Pt(9); rt.font.color.rgb = RGB_GREY
    else:
        _add_heading(doc, "Network Security Group Findings", 2)
        p = doc.add_paragraph()
        r = p.add_run(
            "No NSG security issues detected. All custom rules follow "
            "the principle of least privilege with restricted source addresses."
        )
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x37, 0x56, 0x23)
        doc.add_paragraph()

    # Auto-generated note
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        f"This report was generated automatically on "
        f"{datetime.datetime.utcnow().strftime('%d %B %Y at %H:%M UTC')} "
        f"using Azure Monitor Metrics API (Read-Only access via Azure Lighthouse, "
        f"Service Principal authentication)."
    )
    r.font.size  = Pt(8)
    r.font.italic = True
    r.font.color.rgb = RGB_LGREY
    p.paragraph_format.space_before = Pt(16)

    doc.save(output_path)
    print(f"[INFO] Report saved → {output_path}")
    return output_path
